import os, io, re, time
from pathlib import Path
import streamlit as st

st.set_page_config(page_title="HSE Assistant", layout="wide")

# ───── Secrets ─────
TENANT_ID = st.secrets["TENANT_ID"]
CLIENT_ID = st.secrets["CLIENT_ID"]
CLIENT_SECRET = st.secrets["CLIENT_SECRET"]
MISTRAL_KEY = st.secrets["MISTRAL_API_KEY"]

GRAPH = "https://graph.microsoft.com/v1.0"
TOKEN_URL = f"https://login.microsoftonline.com/{TENANT_ID}/oauth2/v2.0/token"

# ───── UI ─────
st.markdown("""
<style>
body {background:white;}
.bubble{padding:12px;border-radius:10px;max-width:75%;}
.user{background:#eef2f7;margin-left:auto;}
.bot{background:white;border:1px solid #e2e8f0;}
.kw{background:yellow;font-weight:bold;}
</style>
""", unsafe_allow_html=True)

st.title("🦺 HSE Notifications Assistant")

example_q = [
    "PPE requirements for chemical handling",
    "Hot work permit rules",
    "Emergency evacuation steps"
]

st.markdown("**Example Questions:**")
st.write(" • ".join(example_q))

# ───── AUTH ─────
def get_token():
    import requests
    r = requests.post(TOKEN_URL, data={
        "client_id": CLIENT_ID,
        "client_secret": CLIENT_SECRET,
        "scope": "https://graph.microsoft.com/.default",
        "grant_type": "client_credentials"
    })
    return r.json()["access_token"]


def graph_get(url, token):
    import requests
    if not url.startswith("http"):
        url = GRAPH + url
    return requests.get(url, headers={"Authorization": f"Bearer {token}"}).json()


def download(url, token):
    import requests
    return requests.get(url).content


# ───── RECURSIVE FETCH ─────
def fetch_recursive(token, drive_id, folder=""):
    files = []
    url = f"/drives/{drive_id}/root:/{folder}:/children" if folder else f"/drives/{drive_id}/root/children"
    items = graph_get(url, token).get("value", [])

    for item in items:
        if "file" in item:
            if Path(item["name"]).suffix.lower() in [".pdf",".docx"]:
                data = download(item["@microsoft.graph.downloadUrl"], token)
                files.append((item["name"], data, item["@microsoft.graph.downloadUrl"]))
        elif "folder" in item:
            sub = item["name"] if not folder else f"{folder}/{item['name']}"
            files += fetch_recursive(token, drive_id, sub)

    return files


def fetch_files(token, sp_url):
    hostname = sp_url.split("/")[2]
    site_path = "/" + sp_url.split("/", 3)[3]

    site = graph_get(f"/sites/{hostname}:{site_path}", token)
    site_id = site["id"]

    drives = graph_get(f"/sites/{site_id}/drives", token)["value"]
    drive_id = drives[0]["id"]

    return fetch_recursive(token, drive_id)


# ───── PARSER ─────
from langchain_core.documents import Document

def parse_pdf(name,data,link):
    import pdfplumber
    docs=[]
    with pdfplumber.open(io.BytesIO(data)) as pdf:
        for i,p in enumerate(pdf.pages):
            txt=p.extract_text() or ""
            if txt.strip():
                docs.append(Document(
                    page_content=txt,
                    metadata={"file":name,"page":i,"link":link}
                ))
    return docs


def parse_docx(name,data,link):
    import docx
    d=docx.Document(io.BytesIO(data))
    txt="\n".join(p.text for p in d.paragraphs)
    return [Document(page_content=txt,metadata={"file":name,"link":link})]


def parse(name,data,link):
    if name.endswith(".pdf"):
        return parse_pdf(name,data,link)
    if name.endswith(".docx"):
        return parse_docx(name,data,link)
    return []


# ───── INDEX ─────
@st.cache_resource
def build_index(files):
    from langchain_text_splitters import RecursiveCharacterTextSplitter
    from langchain_community.vectorstores import Chroma
    from langchain_community.embeddings import HuggingFaceEmbeddings

    docs=[]
    for n,d,l in files:
        docs+=parse(n,d,l)

    splitter=RecursiveCharacterTextSplitter(chunk_size=800,chunk_overlap=150)
    chunks=splitter.split_documents(docs)

    emb = HuggingFaceEmbeddings(model_name="sentence-transformers/all-MiniLM-L6-v2")
    vs = Chroma.from_documents(chunks, emb)

    return vs, chunks


# ───── HIGHLIGHT ─────
def highlight(text,q):
    words=[w for w in q.split() if len(w)>3]
    for w in words:
        text=re.sub(f"(?i)({w})", r"<span class='kw'>\1</span>", text)
    return text


# ───── RAG ─────
def ask(vs,q):
    from langchain_mistralai import ChatMistralAI
    llm=ChatMistralAI(model="mistral-small-latest",api_key=MISTRAL_KEY)

    docs=vs.similarity_search(q,k=6)

    context="\n\n".join(d.page_content for d in docs)

    prompt=f"""
Use only context.

{context}

Question: {q}
"""

    res=llm.invoke(prompt).content
    return res,docs


# ───── STATE ─────
if "vs" not in st.session_state:
    st.session_state.vs=None
if "chunks" not in st.session_state:
    st.session_state.chunks=[]
if "messages" not in st.session_state:
    st.session_state.messages=[]

# ───── SIDEBAR ─────
with st.sidebar:
    url=st.text_input("SharePoint URL")

    if st.button("Fetch & Index"):
        token=get_token()
        files=fetch_files(token,url)

        vs,chunks=build_index(files)

        st.session_state.vs=vs
        st.session_state.chunks=chunks

        st.success(f"{len(files)} files indexed")

# ───── CHAT ─────
for m in st.session_state.messages:
    cls="user" if m["role"]=="user" else "bot"
    st.markdown(f"<div class='bubble {cls}'>{m['content']}</div>",unsafe_allow_html=True)

# ───── INPUT ─────
q=st.text_input("Ask something")

if st.button("Send") and q and st.session_state.vs:

    st.session_state.messages.append({"role":"user","content":q})

    ans,docs = ask(st.session_state.vs,q)

    # highlight
    ans = highlight(ans,q)

    # sources
    src_html="<br>".join([
        f"📄 <a href='{d.metadata.get('link','#')}' target='_blank'>{d.metadata.get('file')}</a>"
        + (f" p{d.metadata.get('page',0)+1}" if "page" in d.metadata else "")
        for d in docs
    ])

    final=f"{ans}<br><br><b>Sources:</b><br>{src_html}"

    st.session_state.messages.append({"role":"assistant","content":final})

    # chunk viewer
    with st.expander("🔍 View retrieved chunks"):
        for d in docs:
            st.write(d.metadata.get("file"), d.metadata.get("page"))
            st.write(d.page_content[:500])

    st.rerun()
